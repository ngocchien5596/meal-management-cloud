from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

# Read markdown file
md_path = r"g:\Chientest\docs\PLAN-meal-management-tech.md"
with open(md_path, 'r', encoding='utf-8') as f:
    content = f.read()

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

# Title
title = doc.add_heading('KẾ HOẠCH CÔNG NGHỆ – HỆ THỐNG QUẢN LÝ SUẤT ĂN', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Metadata
meta = doc.add_paragraph()
meta.add_run('Dự án: ').bold = True
meta.add_run('Hệ thống Quản lý Suất Ăn (MEAL-2026)\n')
meta.add_run('Ngày tạo: ').bold = True
meta.add_run('01/02/2026\n')
meta.add_run('Mode: ').bold = True
meta.add_run('Self-Hosted (On-Premise)\n')
meta.add_run('Status: ').bold = True
run = meta.add_run('✅ CONFIRMED - Option D')
run.bold = True

doc.add_paragraph('─' * 50)

# Parse sections from markdown
sections = content.split('\n## ')

for section in sections[1:]:
    lines = section.strip().split('\n')
    section_title = lines[0].strip()
    
    # Add section heading
    doc.add_heading(section_title, level=1)
    
    in_table = False
    table_rows = []
    in_code = False
    code_lines = []
    
    for line in lines[1:]:
        line = line.rstrip()
        
        # Handle code blocks
        if line.startswith('```'):
            if in_code:
                if code_lines:
                    code_para = doc.add_paragraph()
                    code_para.paragraph_format.left_indent = Cm(0.5)
                    for cl in code_lines:
                        run = code_para.add_run(cl + '\n')
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        
        if in_code:
            code_lines.append(line)
            continue
        
        # Handle tables
        if line.startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            if '---' not in line:
                cells = [c.strip() for c in line.split('|')[1:-1]]
                table_rows.append(cells)
            continue
        elif in_table:
            if table_rows:
                num_cols = len(table_rows[0])
                num_rows = len(table_rows)
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Table Grid'
                for i, row_data in enumerate(table_rows):
                    for j, cell_data in enumerate(row_data):
                        if j < len(table.rows[i].cells):
                            table.rows[i].cells[j].text = cell_data
                            if i == 0:
                                for para in table.rows[i].cells[j].paragraphs:
                                    for run in para.runs:
                                        run.bold = True
                doc.add_paragraph()
            table_rows = []
            in_table = False
        
        # Handle headings
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('  - '):
            p = doc.add_paragraph(line[4:], style='List Bullet 2')
        elif line.startswith('> '):
            text = line[2:].strip()
            p = doc.add_paragraph()
            clean_text = text.replace('**', '')
            run = p.add_run(clean_text)
            run.italic = True
        elif re.match(r'^\d+\.', line):
            doc.add_paragraph(line, style='List Number')
        elif line.strip():
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*[^*]+\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)

    # Handle remaining table
    if in_table and table_rows:
        num_cols = len(table_rows[0])
        num_rows = len(table_rows)
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        for i, row_data in enumerate(table_rows):
            for j, cell_data in enumerate(row_data):
                if j < len(table.rows[i].cells):
                    table.rows[i].cells[j].text = cell_data
                    if i == 0:
                        for para in table.rows[i].cells[j].paragraphs:
                            for run in para.runs:
                                run.bold = True
        doc.add_paragraph()

# Add footer
doc.add_paragraph()
doc.add_paragraph('─' * 50)
footer = doc.add_paragraph()
footer.add_run('Kế hoạch Công nghệ | Version 1.0 | 01/02/2026 | Option D - Self-Hosted').italic = True
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save
output_path = r"g:\Chientest\docs\KeHoachCongNghe_MealManagement.docx"
doc.save(output_path)
print(f"✅ Word document saved to: {output_path}")
